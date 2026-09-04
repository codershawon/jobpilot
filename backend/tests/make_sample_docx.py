"""Generates backend/sample_cvs/sample_cv.docx for manual/automated testing."""

from pathlib import Path
from docx import Document

OUT = Path(__file__).resolve().parent.parent / "sample_cvs" / "sample_cv.docx"


def build():
    doc = Document()
    doc.add_heading("Farhan Ahmed", level=1)
    doc.add_paragraph("Email: farhan.ahmed@example.com | Phone: +880 1XXXXXXXXX | Dhaka, Bangladesh")

    doc.add_heading("Summary", level=2)
    doc.add_paragraph(
        "Backend-focused software engineer with 4 years of experience building APIs "
        "in Python and Node.js, comfortable working remotely with distributed teams."
    )

    doc.add_heading("Skills", level=2)
    doc.add_paragraph("Python, FastAPI, Django, Node.js, PostgreSQL, Docker, AWS, REST APIs, Git")

    doc.add_heading("Experience", level=2)
    doc.add_paragraph("Backend Engineer — Pathao (Jan 2022 - Present)")
    doc.add_paragraph("- Built and maintained ride-matching microservices handling 2M+ daily requests")
    doc.add_paragraph("- Migrated legacy PHP services to FastAPI, cutting latency by 35%")

    doc.add_paragraph("Software Engineer — Selise Digital (Jun 2020 - Dec 2021)")
    doc.add_paragraph("- Developed internal tooling in Node.js for client reporting dashboards")

    doc.add_heading("Education", level=2)
    doc.add_paragraph("B.Sc. in Computer Science and Engineering — BUET, 2020")

    doc.add_heading("Preferences", level=2)
    doc.add_paragraph("Looking for: Backend Engineer / Software Engineer roles. Open to remote work.")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
